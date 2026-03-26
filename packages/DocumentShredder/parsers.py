"""
DocumentShredder - parsers.py
File-type specific content extraction engines.
"""
import re
from pathlib import Path
from typing import List
from .models import Chunk, ChunkType, ParseResult


# ── helpers ──────────────────────────────────────────────────────────────────

def _count_tokens(text: str) -> int:
    """Rough token estimate: 1 token ≈ 4 chars."""
    try:
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
        return len(enc.encode(text))
    except Exception:
        return max(1, len(text) // 4)


def _make_chunks(paragraphs: List[str], source: str,
                 chunk_type: ChunkType = ChunkType.PROSE,
                 max_tokens: int = 512) -> List[Chunk]:
    chunks, current, cid = [], [], 0
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        current.append(para)
        if _count_tokens(" ".join(current)) >= max_tokens:
            text = "\n\n".join(current)
            chunks.append(Chunk(cid, source, text, chunk_type,
                                tokens=_count_tokens(text)))
            cid += 1
            current = []
    if current:
        text = "\n\n".join(current)
        chunks.append(Chunk(cid, source, text, chunk_type,
                            tokens=_count_tokens(text)))
    return chunks


# ── parsers ──────────────────────────────────────────────────────────────────

def parse_chatgpt_html(filepath: str) -> ParseResult:
    """
    Parse a ChatGPT conversation export (conversations.html).
    Extracts human/assistant turns as clean Q&A chunks.
    """
    from bs4 import BeautifulSoup

    source = Path(filepath).name
    html = Path(filepath).read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "lxml")

    chunks: List[Chunk] = []
    cid = 0

    # ChatGPT export wraps messages in <div data-message-author-role="...">
    messages = soup.find_all(attrs={"data-message-author-role": True})

    if messages:
        for msg in messages:
            role = msg.get("data-message-author-role", "unknown")
            text = msg.get_text(separator="\n", strip=True)
            if not text:
                continue
            ctype = ChunkType.Q_AND_A
            tags = [role, "chatgpt"]
            labeled = f"[{role.upper()}]\n{text}"
            chunks.append(Chunk(cid, source, labeled, ctype,
                                tokens=_count_tokens(labeled), tags=tags))
            cid += 1
    else:
        # Fallback: strip all HTML, split by double newlines
        text = soup.get_text(separator="\n")
        paras = [p for p in text.split("\n\n") if p.strip()]
        chunks = _make_chunks(paras, source, ChunkType.PROSE)

    total = sum(c.tokens for c in chunks)
    return ParseResult(source, chunks, total, "chatgpt_html")


def parse_pdf(filepath: str) -> ParseResult:
    """Extract text from a PDF using pdfplumber; page-aware."""
    import pdfplumber

    source = Path(filepath).name
    chunks: List[Chunk] = []
    cid = 0

    with pdfplumber.open(filepath) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            for table in page.extract_tables():
                rows = [" | ".join(str(c) for c in row if c) for row in table if row]
                md_table = "\n".join(rows)
                chunks.append(Chunk(cid, source, md_table, ChunkType.TABLE,
                                    tokens=_count_tokens(md_table), page=page_num))
                cid += 1
                # Remove table text from prose
                for row in table:
                    for cell in row:
                        text = text.replace(str(cell), "")
            paras = [p.strip() for p in text.split("\n\n") if p.strip()]
            for para in paras:
                chunks.append(Chunk(cid, source, para, ChunkType.PROSE,
                                    tokens=_count_tokens(para), page=page_num))
                cid += 1

    total = sum(c.tokens for c in chunks)
    return ParseResult(source, chunks, total, "pdf")


def parse_docx(filepath: str) -> ParseResult:
    """Extract text and tables from Word documents."""
    from docx import Document

    source = Path(filepath).name
    chunks: List[Chunk] = []
    doc = Document(filepath)
    cid = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        ctype = ChunkType.HEADING if para.style.name.startswith("Heading") else ChunkType.PROSE
        chunks.append(Chunk(cid, source, text, ctype, tokens=_count_tokens(text)))
        cid += 1

    for table in doc.tables:
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        md_table = "\n".join(rows)
        chunks.append(Chunk(cid, source, md_table, ChunkType.TABLE,
                            tokens=_count_tokens(md_table)))
        cid += 1

    total = sum(c.tokens for c in chunks)
    return ParseResult(source, chunks, total, "docx")


def parse_xlsx(filepath: str) -> ParseResult:
    """Convert Excel sheets to Markdown tables."""
    from openpyxl import load_workbook

    source = Path(filepath).name
    chunks: List[Chunk] = []
    wb = load_workbook(filepath, read_only=True, data_only=True)
    cid = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                rows.append(" | ".join(str(c) if c is not None else "" for c in row))
        if rows:
            header = rows[0]
            sep = " | ".join(["---"] * len(rows[0].split("|")))
            md = f"**Sheet: {sheet_name}**\n{header}\n{sep}\n" + "\n".join(rows[1:])
            chunks.append(Chunk(cid, source, md, ChunkType.TABLE,
                                tokens=_count_tokens(md), tags=[sheet_name]))
            cid += 1

    total = sum(c.tokens for c in chunks)
    return ParseResult(source, chunks, total, "xlsx")


def parse_markdown(filepath: str) -> ParseResult:
    """Parse raw Markdown files, splitting by headings."""
    source = Path(filepath).name
    text = Path(filepath).read_text(encoding="utf-8", errors="replace")
    sections = re.split(r"(?m)^#{1,3} ", text)
    chunks = _make_chunks([s.strip() for s in sections if s.strip()], source)
    total = sum(c.tokens for c in chunks)
    return ParseResult(source, chunks, total, "markdown")


def parse_image(filepath: str) -> ParseResult:
    """OCR an image using pytesseract."""
    import pytesseract
    from PIL import Image

    source = Path(filepath).name
    img = Image.open(filepath)
    text = pytesseract.image_to_string(img)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    chunks = _make_chunks(lines, source, ChunkType.PROSE)
    total = sum(c.tokens for c in chunks)
    return ParseResult(source, chunks, total, "ocr_image")
